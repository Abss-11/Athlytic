from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/abhaychaturvedi/Desktop/Athlytic")
OUT = ROOT / "Athlytic_Detailed_Project_Report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 30, 45)
MUTED = RGBColor(90, 100, 115)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=130, bottom=90, end=130) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_inches: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_inches[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths_inches) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def style_run(run, bold=False, italic=False, size=None, color=None) -> None:
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_para(doc, text="", style=None, align=None, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        style_run(r, bold=bold, italic=italic, color=color, size=size)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = BLUE if level <= 2 else DARK_BLUE
        run.bold = True
    return p


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float], header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = INK
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(9.5)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_callout(doc, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    style_run(r, bold=True, color=DARK_BLUE)
    p.add_run(f"\n{text}")
    set_table_width(table, [6.5])
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Athlytic Project Report")
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


def cover_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, "PROJECT REPORT", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=BLUE, size=18)
    add_para(doc, '"Athlytic: Sports-Performance Platform"', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=INK, size=24)
    add_para(doc, "A full-stack web application for athletes, coaches, and fitness creators", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, color=MUTED, size=12)
    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, "SUBMITTED BY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=DARK_BLUE)
    add_para(doc, "Abhay Chaturvedi", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "PROJECT DOMAIN", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=DARK_BLUE)
    add_para(doc, "Sports Technology | Full-Stack Web Development | Product Design | Performance Analytics", align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(4):
        doc.add_paragraph()
    add_para(doc, "Academic Year 2025-26", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para(doc, "Prepared on May 29, 2026", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED)
    page_break(doc)


def front_matter(doc: Document) -> None:
    add_heading(doc, "CERTIFICATE", 1)
    add_para(
        doc,
        'This is to certify that the project report entitled "Athlytic: Sports-Performance Platform" documents the design, development, deployment, and evaluation of a full-stack sports-performance web application. The project covers frontend engineering, backend API development, data modeling, authentication, analytics workflows, Vercel deployment, and a premium UI/UX redesign for athletes and coaches.',
    )
    add_para(doc, "The report has been prepared as a detailed technical and product documentation artifact for review, presentation, and future development planning.")
    doc.add_paragraph()
    add_table(
        doc,
        ["Project", "Prepared By", "Date"],
        [["Athlytic", "Abhay Chaturvedi", "May 29, 2026"]],
        [2.2, 2.2, 2.1],
        header_fill=LIGHT_GRAY,
    )
    page_break(doc)

    add_heading(doc, "DECLARATION", 1)
    add_para(
        doc,
        'I hereby declare that this report on "Athlytic: Sports-Performance Platform" is prepared based on the current repository state, implemented frontend redesign, backend architecture, and Vercel deployment status. The report presents the project objectives, methodology, system design, implementation, testing, deployment, limitations, and future scope in a structured manner.',
    )
    add_para(doc, "The technical observations are based on inspected source files, build results, local verification, and Vercel deployment output.")
    doc.add_paragraph()
    add_para(doc, "Signature: __________________________", bold=True)
    add_para(doc, "Date: ______________________________", bold=True)
    page_break(doc)

    add_heading(doc, "ACKNOWLEDGEMENT", 1)
    add_para(
        doc,
        "I would like to express my sincere gratitude to everyone who supported the development and refinement of this project. The Athlytic platform combines product thinking, frontend design, backend engineering, and deployment practice, and each of these areas required iterative exploration and careful improvement.",
    )
    add_para(
        doc,
        "I also acknowledge the value of modern open-source technologies such as React, Vite, Tailwind CSS, Express, MongoDB, Chart.js, and Vercel, which made it possible to build and deploy the application in a practical and scalable way.",
    )
    add_para(
        doc,
        "Finally, I recognize the importance of continuous feedback and verification. The latest UI/UX redesign was guided by the goal of making Athlytic feel like a premium modern startup product while preserving the working functionality of the existing application.",
    )
    page_break(doc)

    add_heading(doc, "ABSTRACT", 1)
    add_para(
        doc,
        "Athlytic is a sports-performance platform designed to help athletes, coaches, and fitness creators manage training, nutrition, running, sleep, goals, reports, biomarkers, and coach-athlete interactions from a single web-based workspace. The problem addressed by Athlytic is the fragmentation of athlete performance data across multiple tools and manual tracking systems. This fragmentation makes it difficult for athletes to understand progress and for coaches to make timely, evidence-based decisions.",
    )
    add_para(
        doc,
        "The application uses a React and TypeScript frontend, an Express and MongoDB backend, JWT-based authentication, and Vercel deployment. The frontend includes pages for landing, authentication, onboarding, athlete dashboard, coach dashboard, workouts, nutrition, running, goals, reports, integrations, biomarkers, community, and profile management. The backend exposes REST API modules for each major product area and supports protected routes through authentication middleware.",
    )
    add_para(
        doc,
        "A major redesign was completed to improve visual hierarchy, responsiveness, animations, spacing, empty states, loading states, dashboard composition, chart styling, and conversion flow. The redesigned product now presents a more cohesive athlete-tech identity, with a premium glassmorphism-inspired interface, improved navigation, stronger landing page, split-screen authentication flow, polished cards, and better dashboard responsiveness. The project builds successfully and has been deployed to Vercel production.",
    )
    add_para(
        doc,
        "Keywords: Sports Performance, Athlete Dashboard, Coach Portal, React, TypeScript, Express, MongoDB, Vercel, UI/UX Design, Fitness Analytics, Full-Stack Development.",
    )
    page_break(doc)


def table_of_contents(doc: Document) -> None:
    add_heading(doc, "TABLE OF CONTENTS", 1)
    entries = [
        "Chapter 1: Introduction",
        "Chapter 2: Problem Statement and Objectives",
        "Chapter 3: Literature and Product Inspiration Review",
        "Chapter 4: System Requirements",
        "Chapter 5: System Architecture",
        "Chapter 6: Frontend Design and UI/UX Redesign",
        "Chapter 7: Backend Design and API Layer",
        "Chapter 8: Database and Data Model",
        "Chapter 9: Implementation Details",
        "Chapter 10: Testing and Verification",
        "Chapter 11: Vercel Deployment",
        "Chapter 12: Results and Discussion",
        "Chapter 13: Limitations and Future Scope",
        "Chapter 14: Conclusion",
        "References",
        "Appendix A: Route and API Summary",
        "Appendix B: Current Known Issues",
    ]
    for i, entry in enumerate(entries, start=1):
        add_para(doc, f"{i}. {entry}", style="List Number")
    page_break(doc)


def chapter_1(doc: Document) -> None:
    add_heading(doc, "CHAPTER 1: INTRODUCTION", 1)
    add_heading(doc, "1.1 Introduction", 2)
    add_para(
        doc,
        "Athletic performance is influenced by a combination of training load, nutrition, sleep, recovery, running efficiency, strength progression, goal adherence, and coaching feedback. Athletes often track these factors in different applications, spreadsheets, notebooks, or wearable dashboards. As a result, it becomes difficult to view performance holistically and to convert raw activity data into actionable daily decisions.",
    )
    add_para(
        doc,
        "Athlytic is designed as a centralized performance operating system. It brings multiple athlete workflows into a single web application, allowing users to log meals, workouts, sleep, running sessions, biomarker values, goals, and community activity while giving coaches visibility into athlete profiles and notes. The product direction is inspired by modern performance and productivity platforms that combine data clarity with a polished user experience.",
    )
    add_heading(doc, "1.2 Project Overview", 2)
    add_para(
        doc,
        "The project is a full-stack web platform with a React/Vite frontend and an Express backend. The frontend manages user journeys and visual presentation, while the backend provides authentication, protected API routes, data persistence, reporting, and integration endpoints. Vercel is used for deployment, with frontend static output served from `frontend/dist` and API requests routed through serverless entry points.",
    )
    add_heading(doc, "1.3 Motivation", 2)
    add_bullets(
        doc,
        [
            "Athletes need a single workspace to understand training, nutrition, sleep, and goals together.",
            "Coaches need a simple way to view athlete readiness, compliance, and medical notes.",
            "Fitness creators and performance teams need a premium, professional interface that feels trustworthy.",
            "Modern sports products must combine data density with clean hierarchy, mobile responsiveness, and fast interactions.",
        ],
    )
    add_heading(doc, "1.4 Scope", 2)
    add_para(
        doc,
        "The current scope focuses on the web application. The previously tracked mobile Expo application has been removed from the active project direction, simplifying the repository around the web frontend, backend API, and Vercel deployment. The web product includes athlete workflows, coach workflows, dashboard analytics, logging forms, report generation, and product-level UI/UX improvements.",
    )


def chapter_2(doc: Document) -> None:
    add_heading(doc, "CHAPTER 2: PROBLEM STATEMENT AND OBJECTIVES", 1)
    add_heading(doc, "2.1 Problem Statement", 2)
    add_para(
        doc,
        "Athletes and coaches often lack a unified digital environment for tracking performance indicators. Training logs may exist separately from nutrition logs, running data may be separated from recovery notes, and coach feedback may not be connected to daily athlete data. This leads to fragmented decision-making, poor visibility, inconsistent logging, and reduced motivation.",
    )
    add_callout(
        doc,
        "Problem Definition",
        "To design and implement a full-stack sports-performance platform that centralizes athlete tracking, improves coach-athlete visibility, and provides a premium modern user experience without breaking existing functionality.",
    )
    add_heading(doc, "2.2 Project Objectives", 2)
    add_numbered(
        doc,
        [
            "Build a responsive web platform for athlete and coach workflows.",
            "Provide authentication and role-based navigation for athletes and coaches.",
            "Support logging and analysis of workouts, nutrition, running, sleep, biomarkers, and goals.",
            "Create dashboard views that summarize daily performance and trends.",
            "Generate weekly reports and support PDF report output through the backend.",
            "Improve the UI/UX to feel like a premium athlete-tech startup product.",
            "Deploy the application successfully to Vercel with a reliable build configuration.",
        ],
    )
    add_heading(doc, "2.3 Success Criteria", 2)
    add_table(
        doc,
        ["Criterion", "Expected Outcome", "Current Status"],
        [
            ["Build", "Frontend production build completes without TypeScript or Vite errors.", "Passed locally and on Vercel."],
            ["Deployment", "Vercel production deployment is live.", "Production URL is active."],
            ["UX", "Core flows feel modern, consistent, and responsive.", "Major redesign completed."],
            ["Functionality", "Existing routes and API calls remain intact.", "No functional breaks found during verification."],
            ["Maintainability", "Shared components and styling are reusable.", "Improved through shared card, button, input, chart, and shell updates."],
        ],
        [1.3, 3.2, 2.0],
    )


def chapter_3(doc: Document) -> None:
    add_heading(doc, "CHAPTER 3: LITERATURE AND PRODUCT INSPIRATION REVIEW", 1)
    add_para(
        doc,
        "The redesign and product direction were influenced by established sports, productivity, and modern SaaS applications. Rather than copying any single product, Athlytic combines patterns from performance tracking, clean productivity dashboards, and polished startup interfaces.",
    )
    add_table(
        doc,
        ["Reference Product", "Relevant Strength", "Application to Athlytic"],
        [
            ["Nike Training Club", "Motivational training experience and strong athletic tone.", "Used as inspiration for athlete-focused energy and action-oriented pages."],
            ["Strava", "Activity history, progress visibility, and social motivation.", "Informs running history, workout logs, and community flow."],
            ["Linear", "Clean hierarchy, subtle motion, and polished SaaS design.", "Influences navigation, spacing, panel treatment, and typography."],
            ["Notion", "Flexible information architecture and approachable workspace feel.", "Informs dashboard organization and modular content cards."],
            ["Apple Fitness", "Premium wellness aesthetic and strong visual clarity.", "Informs clean metric presentation and health-focused interface tone."],
            ["Whoop", "Recovery and readiness storytelling.", "Informs readiness, sleep, and performance score framing."],
            ["Vercel", "Minimal developer-product polish and deployment confidence.", "Informs brand restraint, button styling, and product presentation."],
        ],
        [1.6, 2.3, 2.6],
    )
    add_heading(doc, "3.1 Design Principles Derived", 2)
    add_bullets(
        doc,
        [
            "Prioritize the next useful action over raw data density.",
            "Use premium but restrained surfaces, avoiding visual clutter.",
            "Give empty states a helpful purpose instead of leaving blank space.",
            "Make responsive behavior a first-class requirement.",
            "Use motion and hover feedback to make the product feel alive without hurting performance.",
            "Maintain consistent component language across all pages.",
        ],
    )


def chapter_4(doc: Document) -> None:
    add_heading(doc, "CHAPTER 4: SYSTEM REQUIREMENTS", 1)
    add_heading(doc, "4.1 Functional Requirements", 2)
    add_table(
        doc,
        ["Module", "Functional Requirement"],
        [
            ["Authentication", "Users can register, log in, and access role-specific flows using JWT authentication."],
            ["Profile", "Athletes can maintain sport, body metrics, activity level, goals, diet preferences, allergies, and medical context."],
            ["Dashboard", "Athletes can view performance score, readiness, nutrition, sleep, workouts, running, goals, and AI insights."],
            ["Workout", "Users can create, edit, delete, and restore workout sessions with exercises and set logs."],
            ["Nutrition", "Users can log meals, track macros, compare today with yesterday, and receive recommendations."],
            ["Running", "Users can track distance, duration, pace, VO2 Max, and personal running history."],
            ["Coach Portal", "Coaches can view athletes and update coach notes or medical directives."],
            ["Reports", "Users can view weekly reports and download PDF reports."],
            ["Integrations", "Users can trigger provider sync flows for connected fitness sources."],
        ],
        [1.4, 5.1],
    )
    add_heading(doc, "4.2 Non-Functional Requirements", 2)
    add_bullets(
        doc,
        [
            "The application should build reliably in local and Vercel environments.",
            "The UI should be responsive across desktop, tablet, and mobile screen sizes.",
            "The interface should remain accessible through semantic navigation and visible focus states.",
            "The codebase should maintain reusable frontend components.",
            "The backend should provide structured API responses and handle errors predictably.",
            "The deployment path should avoid dependency on the removed mobile application.",
        ],
    )
    add_heading(doc, "4.3 Hardware and Software Requirements", 2)
    add_table(
        doc,
        ["Category", "Technology / Tool", "Purpose"],
        [
            ["Frontend Runtime", "React 19 + Vite", "Build and run the web interface."],
            ["Language", "TypeScript", "Provide typed frontend development."],
            ["Styling", "Tailwind CSS", "Implement responsive UI and design tokens."],
            ["Backend Runtime", "Node.js", "Run Express API server."],
            ["Backend Framework", "Express", "Expose REST API routes."],
            ["Database", "MongoDB + Mongoose", "Persist users, logs, goals, reports, and related data."],
            ["Charts", "Chart.js + React Chart.js 2", "Render dashboard and analytics charts."],
            ["Deployment", "Vercel", "Build and host frontend and serverless API."],
        ],
        [1.5, 2.0, 3.0],
    )


def chapter_5(doc: Document) -> None:
    add_heading(doc, "CHAPTER 5: SYSTEM ARCHITECTURE", 1)
    add_para(
        doc,
        "Athlytic follows a modular full-stack architecture. The frontend is responsible for routing, layout, forms, dashboard composition, charts, and user interaction. The backend is responsible for authentication, data validation, business logic, persistence, reporting, and integration routes. Vercel acts as the hosting and routing layer.",
    )
    add_heading(doc, "5.1 High-Level Architecture", 2)
    add_table(
        doc,
        ["Layer", "Main Responsibility", "Important Files"],
        [
            ["Presentation Layer", "User interface, pages, navigation, forms, cards, charts, and design system.", "frontend/src/pages, frontend/src/components, frontend/src/index.css"],
            ["Client API Layer", "Axios client and typed API helpers for frontend-backend communication.", "frontend/src/api/api.ts"],
            ["Backend API Layer", "Express app, route registration, error handling, and module APIs.", "backend/src/app.js, backend/src/routes"],
            ["Business Logic Layer", "Macro planning, daily metrics, persistence utilities, token generation, validation.", "backend/src/utils"],
            ["Data Layer", "Mongoose models for users, workouts, nutrition, sleep, goals, biomarkers, and community.", "backend/src/models"],
            ["Deployment Layer", "Vercel build, frontend output, API rewrites, serverless entry point.", "vercel.json, api/index.js"],
        ],
        [1.45, 2.85, 2.2],
    )
    add_heading(doc, "5.2 Request Flow", 2)
    add_numbered(
        doc,
        [
            "The user opens the Athlytic frontend through the Vercel production URL.",
            "React Router resolves the requested route and renders the correct page or redirects unauthenticated users to login.",
            "The frontend sends API requests through Axios using `/api` as the base path.",
            "Vercel rewrites `/api` requests to the serverless API entry point.",
            "The Express app routes the request to the appropriate module.",
            "Protected routes verify the JWT token through authentication middleware.",
            "The backend reads or writes data using Mongoose models and returns JSON responses.",
            "The frontend updates local state and re-renders dashboards, forms, charts, or feedback states.",
        ],
    )
    add_heading(doc, "5.3 Repository Organization", 2)
    add_para(
        doc,
        "The repository uses root-level npm scripts to delegate frontend and backend commands. The active workspaces are `frontend` and `backend`. The removed `mobile` folder is no longer part of the Vercel build path. The `api` folder provides Vercel serverless routing into the backend application.",
    )


def chapter_6(doc: Document) -> None:
    add_heading(doc, "CHAPTER 6: FRONTEND DESIGN AND UI/UX REDESIGN", 1)
    add_heading(doc, "6.1 Previous UI Limitations", 2)
    add_para(
        doc,
        "Before the redesign, the application had working product flows but the interface felt basic and disconnected. Page-level layouts, cards, spacing, hierarchy, and visual feedback were not consistent enough for a premium sports-performance product. The landing page did not communicate a strong brand identity, and the authenticated dashboard needed more polish and responsiveness.",
    )
    add_heading(doc, "6.2 Redesign Goals", 2)
    add_bullets(
        doc,
        [
            "Make Athlytic feel like a premium modern startup product.",
            "Create a stronger athlete-tech visual identity.",
            "Improve the landing page conversion flow.",
            "Improve role-based authentication screens.",
            "Make dashboard widgets feel alive and data-rich.",
            "Improve forms, cards, chart styling, spacing, and mobile behavior.",
            "Add tasteful transitions, hover states, and skeleton loading patterns.",
        ],
    )
    add_heading(doc, "6.3 Design System Improvements", 2)
    add_table(
        doc,
        ["Component / Area", "Improvement"],
        [
            ["CSS tokens", "Refined light/dark variables for backgrounds, surfaces, borders, primary color, accent color, success, warning, danger, shadows, and panel glow."],
            ["Cards", "Improved glass panel treatment, hover lift, consistent padding, subtle surface gradients, and stronger borders."],
            ["Buttons", "Updated primary, secondary, and ghost variants with better hover, active, disabled, and focus states."],
            ["Inputs", "Improved hover/focus behavior, rounded form style, and visible focus ring."],
            ["Badges", "Added accent dot and tighter uppercase visual language."],
            ["Charts", "Improved card sizing, tooltip styling, line tension, bar radius, and skeleton empty states."],
            ["Page headers", "Improved title hierarchy, spacing, and badge placement."],
            ["App shell", "Redesigned navigation with brand mark, horizontal mobile nav, active states, alerts dropdown, and improved account panel."],
        ],
        [1.8, 4.7],
    )
    add_heading(doc, "6.4 Page-Level Improvements", 2)
    add_table(
        doc,
        ["Page", "Redesign Summary"],
        [
            ["Landing", "Rebuilt as a premium conversion page with dark hero, proof metrics, live cockpit card, workflow, feature cards, pricing, testimonials, and CTA footer."],
            ["Login", "Converted into split-screen experience with role selection and contextual performance messaging."],
            ["Register", "Improved onboarding expectation and role-based account creation."],
            ["Onboarding", "Added step-based layout with progress indicator and clearer profile setup path."],
            ["Dashboard", "Improved stat grid, loading skeletons, daily plan cards, empty states, chart cards, and responsive breakpoints."],
            ["Workout", "Improved session builder, exercise cards, set logs, analytics cards, and recent sessions."],
            ["Nutrition", "Improved macro tracker, meal log cards, guidance states, and visual hierarchy."],
            ["Coach", "Improved athlete overview, notes area, action queue, and comparison metrics."],
        ],
        [1.45, 5.05],
    )
    add_heading(doc, "6.5 Responsiveness and Motion", 2)
    add_para(
        doc,
        "The redesign introduces an `app-page` entrance animation, consistent interactive card transitions, skeleton loaders, hover lift states, and responsive grid changes. The dashboard stat grid was adjusted so cards do not become cramped on laptop-width screens. Motion is intentionally lightweight and CSS-based to avoid adding heavy runtime dependencies.",
    )


def chapter_7(doc: Document) -> None:
    add_heading(doc, "CHAPTER 7: BACKEND DESIGN AND API LAYER", 1)
    add_para(
        doc,
        "The backend is an Express application organized around route modules. It uses middleware for JSON parsing, CORS, authentication, and error handling. The backend supports both local full-stack serving and Vercel serverless deployment.",
    )
    add_heading(doc, "7.1 API Modules", 2)
    add_table(
        doc,
        ["API Module", "Base Route", "Primary Responsibility"],
        [
            ["Authentication", "/api/auth", "Register and login users, generate JWT tokens."],
            ["Profile", "/api/profile", "Read and update user profile and macro plan data."],
            ["Dashboard", "/api/dashboard", "Serve athlete dashboard, coach summary, and notifications."],
            ["Nutrition", "/api/nutrition", "CRUD meal logs and macro summaries."],
            ["Running", "/api/running", "CRUD running sessions."],
            ["Sleep", "/api/sleep", "CRUD sleep logs and summaries."],
            ["Workouts", "/api/workouts", "CRUD workouts and workout summaries."],
            ["Goals", "/api/goals", "Goal listing and creation."],
            ["Reports", "/api/reports", "Weekly reports and PDF download."],
            ["Coach", "/api/coach", "Coach athlete list and coach notes update."],
            ["Biomarkers", "/api/biomarkers", "Biomarker log listing, creation, and deletion."],
        ],
        [1.45, 1.45, 3.6],
    )
    add_heading(doc, "7.2 Authentication and Protection", 2)
    add_para(
        doc,
        "Several route modules use the `protect` middleware to require a valid JWT token. The frontend stores the token in local storage and attaches it to requests through the Axios interceptor. Protected workflows include profile, dashboard, nutrition, running, sleep, workouts, reports, coach, and biomarkers.",
    )
    add_heading(doc, "7.3 Error Handling", 2)
    add_para(
        doc,
        "The Express app includes a final error-handling middleware that logs the error and returns a JSON response with a status code and message. This provides a consistent fallback for unexpected server errors while allowing route modules to return specific validation or authorization messages.",
    )


def chapter_8(doc: Document) -> None:
    add_heading(doc, "CHAPTER 8: DATABASE AND DATA MODEL", 1)
    add_para(
        doc,
        "The backend uses Mongoose models to represent the key entities in the Athlytic domain. These models support the central product workflows: users, workouts, nutrition, running, sleep, goals, biomarkers, achievements, and community posts.",
    )
    add_table(
        doc,
        ["Model", "Purpose"],
        [
            ["User", "Stores user identity, role, profile metrics, preferences, and coach-relevant notes."],
            ["Workout", "Stores workout sessions, exercises, set logs, duration, intensity, and derived strength metrics."],
            ["NutritionLog", "Stores meals, calories, protein, carbs, fats, and water intake."],
            ["RunningData", "Stores running distance, duration, pace, VO2 Max, and personal-record metadata."],
            ["SleepLog", "Stores daily sleep hours and notes."],
            ["Goal", "Stores athlete goals with category, target value, current value, unit, and deadline."],
            ["BiomarkerLog", "Stores biomarker entries such as health and performance measurements."],
            ["CommunityPost", "Stores social/community updates and engagement data."],
            ["Achievement", "Stores achievement or milestone information."],
        ],
        [1.7, 4.8],
    )
    add_heading(doc, "8.1 Persistence Strategy", 2)
    add_para(
        doc,
        "The backend includes persistence utilities and sample data support. When MongoDB is available, data is persisted through Mongoose. During local fallback conditions, mock in-memory behavior can support development and testing. This makes the project easier to run while retaining a real database path for production-style usage.",
    )


def chapter_9(doc: Document) -> None:
    add_heading(doc, "CHAPTER 9: IMPLEMENTATION DETAILS", 1)
    add_heading(doc, "9.1 Frontend Routing", 2)
    add_para(
        doc,
        "The frontend uses React Router with lazy-loaded page components. Public routes include landing, login, and signup. Authenticated routes are nested inside the `AppShell`, which provides navigation, alerts, account actions, theme toggle, and shared page chrome.",
    )
    add_table(
        doc,
        ["Route", "Page"],
        [
            ["/", "Landing page"],
            ["/login", "Login page"],
            ["/signup", "Registration page"],
            ["/dashboard", "Athlete dashboard"],
            ["/onboarding", "Athlete onboarding"],
            ["/coach", "Coach dashboard"],
            ["/nutrition", "Nutrition tracker"],
            ["/workouts", "Workout tracker"],
            ["/running", "Running tracker"],
            ["/goals", "Goals page"],
            ["/community", "Community page"],
            ["/profile", "Profile page"],
            ["/reports", "Reports page"],
            ["/integrations", "Integrations page"],
            ["/biomarkers", "Biomarkers page"],
        ],
        [1.5, 5.0],
    )
    add_heading(doc, "9.2 Frontend API Client", 2)
    add_para(
        doc,
        "The frontend API client is centralized in `frontend/src/api/api.ts`. It creates a shared Axios instance, uses `/api` as the base URL by default, and attaches the JWT token from local storage to authenticated requests. Each product module has a grouped API helper object, improving readability and reducing duplicate request code.",
    )
    add_heading(doc, "9.3 Backend Route Registration", 2)
    add_para(
        doc,
        "The backend registers each major product domain as a route module under `/api`. This keeps route files focused and allows independent development of authentication, profile, dashboard, nutrition, running, sleep, workout, goal, community, report, integration, coach, and biomarker modules.",
    )
    add_heading(doc, "9.4 Report Generation", 2)
    add_para(
        doc,
        "The reports module supports weekly report retrieval and PDF generation. PDFKit is included in the backend dependency tree and is used to generate downloadable report output. This feature is important because coaches and athletes often need shareable summaries beyond the live dashboard.",
    )


def chapter_10(doc: Document) -> None:
    add_heading(doc, "CHAPTER 10: TESTING AND VERIFICATION", 1)
    add_heading(doc, "10.1 Build Verification", 2)
    add_para(
        doc,
        "The project was verified through a production build using the root command `npm run build`, which delegates to the frontend build script. The build completed successfully after type and CSS utility issues introduced during redesign were corrected.",
    )
    add_heading(doc, "10.2 Browser Verification", 2)
    add_bullets(
        doc,
        [
            "Landing page was opened locally and visually inspected.",
            "Hero contrast issue was identified and corrected.",
            "Login page was opened locally and visually inspected.",
            "Unauthenticated dashboard access redirected correctly to login.",
            "Authenticated dashboard entry was verified locally.",
            "Production landing page was opened and visually verified after Vercel deployment.",
        ],
    )
    add_heading(doc, "10.3 Lint and Known Quality Signals", 2)
    add_para(
        doc,
        "The frontend lint command currently reports remaining pre-existing warnings/errors in React context and biomarkers code. These are not blocking the production build but should be addressed to improve maintainability and development quality.",
    )
    add_table(
        doc,
        ["Check", "Result", "Comment"],
        [
            ["Local frontend build", "Passed", "TypeScript and Vite build completed successfully."],
            ["Vercel production build", "Passed", "Vercel completed install, build, and deploy."],
            ["Landing render", "Passed", "Production landing page renders redesigned UI."],
            ["Protected route redirect", "Passed locally", "Dashboard redirects to login when unauthenticated."],
            ["Frontend lint", "Open issues", "Context export warnings and one biomarkers hook dependency warning remain."],
            ["npm audit", "Open issues", "Vercel reported backend and frontend vulnerabilities."],
        ],
        [1.7, 1.4, 3.4],
    )


def chapter_11(doc: Document) -> None:
    add_heading(doc, "CHAPTER 11: VERCEL DEPLOYMENT", 1)
    add_para(
        doc,
        "The project is linked to Vercel and configured for deployment through `vercel.json`. The production deployment completed successfully and was aliased to the main production URL.",
    )
    add_heading(doc, "11.1 Vercel Configuration", 2)
    add_table(
        doc,
        ["Vercel Setting", "Value"],
        [
            ["Install command", "npm --prefix backend install && npm --prefix frontend install"],
            ["Build command", "npm --prefix frontend run build"],
            ["Output directory", "frontend/dist"],
            ["API rewrite", "/api and /api/:path* route to /api/index"],
            ["SPA rewrite", "Non-API routes route to /index.html"],
        ],
        [2.1, 4.4],
    )
    add_heading(doc, "11.2 Deployment URLs", 2)
    add_table(
        doc,
        ["Environment", "URL"],
        [
            ["Production alias", "https://athlytic-pink.vercel.app"],
            ["Production deployment", "https://athlytic-hdzh2g13z-abhaychaturvedi485-2309s-projects.vercel.app"],
            ["Preview deployment", "https://athlytic-6tlyciic4-abhaychaturvedi485-2309s-projects.vercel.app"],
        ],
        [2.0, 4.5],
    )
    add_heading(doc, "11.3 Deployment Result", 2)
    add_para(
        doc,
        "The Vercel production deployment reached the READY state. The production page was opened and verified to render the redesigned landing experience. Vercel install output reported dependency audit warnings, which should be reviewed as a follow-up security task.",
    )


def chapter_12(doc: Document) -> None:
    add_heading(doc, "CHAPTER 12: RESULTS AND DISCUSSION", 1)
    add_heading(doc, "12.1 Product Outcome", 2)
    add_para(
        doc,
        "The latest version of Athlytic presents a stronger and more cohesive product experience. The landing page now clearly communicates the product category, value proposition, workflow, proof metrics, pricing structure, and conversion actions. The authenticated app shell has a clearer brand identity and better navigation. The dashboard and tracking pages now share a more consistent component language.",
    )
    add_heading(doc, "12.2 Technical Outcome", 2)
    add_para(
        doc,
        "The technical architecture remains stable after the redesign. The frontend build passes locally and on Vercel. The Vercel deployment path is aligned with the current web-focused repository structure. The removed mobile app is no longer required for deployment.",
    )
    add_heading(doc, "12.3 UX Outcome", 2)
    add_table(
        doc,
        ["Before", "After"],
        [
            ["Basic landing page with weaker conversion hierarchy.", "Premium landing page with strong hero, proof metrics, workflow, pricing, testimonials, and CTA."],
            ["Disconnected card and layout patterns.", "Consistent glass panels, surface tiles, badges, buttons, inputs, and page headers."],
            ["Dashboard stat cards could feel cramped.", "Responsive dashboard grid adjusted for laptop and wider screens."],
            ["Auth flow felt utilitarian.", "Split-screen login/register flow with role selection and stronger product framing."],
            ["Empty states were plain.", "Empty states now communicate next action and use dashed/skeleton visual patterns."],
        ],
        [3.2, 3.3],
    )


def chapter_13(doc: Document) -> None:
    add_heading(doc, "CHAPTER 13: LIMITATIONS AND FUTURE SCOPE", 1)
    add_heading(doc, "13.1 Current Limitations", 2)
    add_bullets(
        doc,
        [
            "Production authenticated API behavior still needs complete verification against production environment variables and database state.",
            "Frontend lint still has remaining context and hook warnings.",
            "Dependency audit warnings need review and remediation.",
            "Some secondary pages, including reports, integrations, biomarkers, and community, can receive deeper visual polish.",
            "Automated end-to-end tests are not yet established.",
            "The current product uses REST APIs and local client state rather than a more advanced query cache layer.",
        ],
    )
    add_heading(doc, "13.2 Future Scope", 2)
    add_numbered(
        doc,
        [
            "Add end-to-end test coverage for login, dashboard load, meal logging, workout creation, running session creation, and coach notes.",
            "Improve production observability with Vercel logs, analytics, error tracking, and health checks.",
            "Add richer charts for readiness, load progression, macro adherence, and training volume.",
            "Introduce coach-athlete messaging or feedback workflows.",
            "Add wearable integrations for automatic running, sleep, heart-rate, and recovery data.",
            "Improve AI recommendation logic with explainable insights and personalized coaching suggestions.",
            "Add team management features for coaches working with multiple athletes or groups.",
            "Add mobile-responsive progressive web app enhancements now that the native mobile folder has been removed.",
        ],
    )


def chapter_14(doc: Document) -> None:
    add_heading(doc, "CHAPTER 14: CONCLUSION", 1)
    add_para(
        doc,
        "Athlytic is a promising full-stack sports-performance platform that brings together athlete tracking, coach visibility, analytics, reporting, and modern product design. The latest UI/UX redesign moved the product from a functional application toward a premium startup-quality experience. The frontend now has stronger visual hierarchy, a clearer conversion path, improved navigation, better card systems, and more polished dashboard and tracking flows.",
    )
    add_para(
        doc,
        "From a technical perspective, the project uses a practical architecture based on React, TypeScript, Vite, Express, MongoDB, and Vercel. The build and deployment pipeline is functioning, and the production URL is live. The next major priorities are production API verification, lint cleanup, dependency security review, test automation, and continued polish of secondary pages.",
    )
    add_para(
        doc,
        "Overall, Athlytic demonstrates how a focused full-stack application can combine performance tracking and modern UI/UX to support athletes, coaches, and fitness creators in making better training and recovery decisions.",
    )


def references_and_appendix(doc: Document) -> None:
    page_break(doc)
    add_heading(doc, "REFERENCES", 1)
    add_numbered(
        doc,
        [
            "React Documentation. Component-based frontend development and React Router integration patterns.",
            "Vite Documentation. Frontend build tooling and production bundling.",
            "Tailwind CSS Documentation. Utility-first styling and responsive design.",
            "Express Documentation. Routing, middleware, and API server design.",
            "MongoDB and Mongoose Documentation. Schema modeling and data persistence.",
            "Vercel Documentation. Static frontend deployment, serverless functions, rewrites, and production deployment.",
            "Chart.js Documentation. Chart rendering, tooltips, scales, and visual configuration.",
            "Product design references: Nike Training Club, Strava, Linear, Notion, Apple Fitness, Vercel, Airbnb, Whoop, and Duolingo animation patterns.",
        ],
    )
    page_break(doc)
    add_heading(doc, "APPENDIX A: ROUTE AND API SUMMARY", 1)
    add_table(
        doc,
        ["Frontend Area", "Route", "Backend API Used"],
        [
            ["Landing", "/", "None"],
            ["Login", "/login", "/api/auth/login, /api/profile/me"],
            ["Signup", "/signup", "/api/auth/register"],
            ["Dashboard", "/dashboard", "/api/dashboard/athlete, /api/sleep, /api/dashboard/notifications"],
            ["Nutrition", "/nutrition", "/api/nutrition, /api/nutrition/summary"],
            ["Workouts", "/workouts", "/api/workouts, /api/workouts/summary"],
            ["Running", "/running", "/api/running"],
            ["Goals", "/goals", "/api/goals"],
            ["Coach", "/coach", "/api/dashboard/coach, /api/coach/athletes"],
            ["Reports", "/reports", "/api/reports/weekly, /api/reports/weekly/pdf"],
            ["Biomarkers", "/biomarkers", "/api/biomarkers"],
        ],
        [1.6, 1.35, 3.55],
    )
    add_heading(doc, "APPENDIX B: CURRENT KNOWN ISSUES", 1)
    add_table(
        doc,
        ["Issue", "Impact", "Suggested Resolution"],
        [
            ["React context lint warnings", "Developer quality issue; build still passes.", "Split non-component exports into separate files or adjust architecture."],
            ["AuthContext setState-in-effect warning", "Potential render cascade concern.", "Initialize state lazily from localStorage or refactor effect logic."],
            ["Biomarkers hook dependency warning", "Potential stale effect dependency.", "Wrap fetch function in useCallback or include it in the dependency array."],
            ["npm audit vulnerabilities", "Security and dependency health concern.", "Run npm audit and patch packages safely."],
            ["Production API verification incomplete", "Authenticated production flows need confirmation.", "Verify production env vars, database, login, and API routes."],
            ["Mixed git changes", "Commit history could become unclear.", "Separate mobile removal, UI redesign, and dependency changes into separate commits."],
        ],
        [1.8, 2.0, 2.7],
    )


def main() -> None:
    doc = Document()
    configure_document(doc)
    cover_page(doc)
    front_matter(doc)
    table_of_contents(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)
    chapter_7(doc)
    chapter_8(doc)
    chapter_9(doc)
    chapter_10(doc)
    chapter_11(doc)
    chapter_12(doc)
    chapter_13(doc)
    chapter_14(doc)
    references_and_appendix(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
